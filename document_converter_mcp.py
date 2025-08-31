#!/usr/bin/env python3
"""
统一文档转换 MCP 服务器
支持 PDF、Word、Markdown、HTML 等格式之间的相互转换
"""

import os
import sys
import json
import time
import logging
import traceback
from pathlib import Path
from typing import Optional, Dict, Any, List
from functools import wraps

# FastMCP imports
from fastmcp import FastMCP

# 文档处理库
try:
    from markitdown import MarkItDown
    MARKITDOWN_AVAILABLE = True
except ImportError:
    MARKITDOWN_AVAILABLE = False

try:
    import pypandoc
    PYPANDOC_AVAILABLE = True
except ImportError:
    PYPANDOC_AVAILABLE = False

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import markdown
    from bs4 import BeautifulSoup
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('document_converter.log')
    ]
)
logger = logging.getLogger(__name__)

# 服务器配置
class ServerConfig:
    def __init__(self):
        self.max_file_size = int(os.getenv('MAX_FILE_SIZE', 100 * 1024 * 1024))  # 100MB
        self.temp_dir = os.getenv('TEMP_DIR', 'temp')
        self.log_level = os.getenv('LOG_LEVEL', 'INFO')
        self.enable_pandoc = os.getenv('ENABLE_PANDOC', 'true').lower() == 'true'
        
    def validate_file_size(self, file_path: str) -> bool:
        """验证文件大小"""
        try:
            file_size = os.path.getsize(file_path)
            return file_size <= self.max_file_size
        except OSError:
            return False

config = ServerConfig()

# 性能监控装饰器
def performance_monitor(func):
    @wraps(func)
    def wrapper(*args, **kwargs):
        start_time = time.time()
        try:
            result = func(*args, **kwargs)
            execution_time = time.time() - start_time
            logger.info(f"{func.__name__} 执行成功，耗时: {execution_time:.3f}秒")
            return result
        except Exception as e:
            execution_time = time.time() - start_time
            logger.error(f"{func.__name__} 执行失败，耗时: {execution_time:.3f}秒，错误: {str(e)}")
            logger.error(f"错误详情: {traceback.format_exc()}")
            raise
    return wrapper

# 创建 FastMCP 应用
mcp = FastMCP("DocumentConverter")

@mcp.tool()
@performance_monitor
def convert_document(input_path: str, target_format: str, output_path: str = "") -> str:
    """
    转换文档格式
    
    Args:
        input_path: 输入文件路径
        target_format: 目标格式 (pdf, docx, markdown, html, txt)
        output_path: 输出文件路径（可选）
    
    Returns:
        转换结果信息
    """
    try:
        # 验证输入文件
        if not os.path.exists(input_path):
            return f"错误: 输入文件不存在: {input_path}"
        
        if not config.validate_file_size(input_path):
            return f"错误: 文件大小超过限制 ({config.max_file_size / 1024 / 1024:.1f}MB)"
        
        # 生成输出路径
        if not output_path:
            input_file = Path(input_path)
            output_path = str(input_file.parent / f"{input_file.stem}.{target_format}")
        
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # 获取输入文件格式
        input_format = Path(input_path).suffix.lower().lstrip('.')
        
        # 执行转换
        if target_format.lower() == 'txt' or target_format.lower() == 'markdown':
            return _convert_to_text(input_path, output_path, target_format)
        elif target_format.lower() == 'html':
            return _convert_to_html(input_path, output_path)
        elif target_format.lower() == 'docx':
            return _convert_to_docx(input_path, output_path)
        elif target_format.lower() == 'pdf':
            return _convert_to_pdf(input_path, output_path)
        else:
            return f"错误: 不支持的目标格式: {target_format}"
            
    except Exception as e:
        return f"转换失败: {str(e)}"

def _convert_to_text(input_path: str, output_path: str, format_type: str) -> str:
    """转换为文本格式"""
    input_ext = Path(input_path).suffix.lower()
    
    try:
        # 使用 MarkItDown 处理各种格式
        if MARKITDOWN_AVAILABLE:
            md = MarkItDown()
            result = md.convert(input_path)
            content = result.text_content
            
            if format_type == 'markdown':
                # 保持 Markdown 格式
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
            else:
                # 转换为纯文本
                if MARKDOWN_AVAILABLE:
                    from bs4 import BeautifulSoup
                    import markdown
                    html = markdown.markdown(content)
                    soup = BeautifulSoup(html, 'html.parser')
                    text = soup.get_text()
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(text)
                else:
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(content)
            
            return f"转换成功: {output_path}"
        
        # 备用方案：直接读取文本文件
        elif input_ext in ['.txt', '.md']:
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(content)
            return f"转换成功: {output_path}"
        
        else:
            return "错误: 缺少必要的转换库"
            
    except Exception as e:
        return f"文本转换失败: {str(e)}"

def _convert_to_html(input_path: str, output_path: str) -> str:
    """转换为HTML格式"""
    input_ext = Path(input_path).suffix.lower()
    
    try:
        if input_ext == '.md' and MARKDOWN_AVAILABLE:
            # Markdown 转 HTML
            with open(input_path, 'r', encoding='utf-8') as f:
                md_content = f.read()
            html_content = markdown.markdown(md_content)
            
            # 添加基本的HTML结构
            full_html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Converted Document</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        pre {{ background-color: #f4f4f4; padding: 10px; border-radius: 5px; }}
        code {{ background-color: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(full_html)
            return f"转换成功: {output_path}"
        
        # 使用 MarkItDown 处理其他格式
        elif MARKITDOWN_AVAILABLE:
            md = MarkItDown()
            result = md.convert(input_path)
            md_content = result.text_content
            
            if MARKDOWN_AVAILABLE:
                html_content = markdown.markdown(md_content)
                full_html = f"""
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Converted Document</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; }}
        pre {{ background-color: #f4f4f4; padding: 10px; border-radius: 5px; }}
        code {{ background-color: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>
"""
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(full_html)
                return f"转换成功: {output_path}"
        
        return "错误: 缺少必要的转换库"
        
    except Exception as e:
        return f"HTML转换失败: {str(e)}"

def _convert_to_docx(input_path: str, output_path: str) -> str:
    """转换为DOCX格式"""
    if not PYTHON_DOCX_AVAILABLE:
        return "错误: 缺少 python-docx 库"
    
    try:
        # 先转换为文本
        if MARKITDOWN_AVAILABLE:
            md = MarkItDown()
            result = md.convert(input_path)
            content = result.text_content
        else:
            with open(input_path, 'r', encoding='utf-8') as f:
                content = f.read()
        
        # 创建Word文档
        doc = Document()
        
        # 按行添加内容
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
            else:
                doc.add_paragraph()
        
        doc.save(output_path)
        return f"转换成功: {output_path}"
        
    except Exception as e:
        return f"DOCX转换失败: {str(e)}"

def _convert_to_pdf(input_path: str, output_path: str) -> str:
    """转换为PDF格式"""
    if PYPANDOC_AVAILABLE and config.enable_pandoc:
        try:
            pypandoc.convert_file(input_path, 'pdf', outputfile=output_path)
            return f"转换成功: {output_path}"
        except Exception as e:
            return f"PDF转换失败: {str(e)}"
    else:
        return "错误: PDF转换需要 pypandoc 和 pandoc"

@mcp.tool()
def list_supported_formats() -> str:
    """
    列出支持的文档格式
    
    Returns:
        支持的格式列表
    """
    formats = {
        "输入格式": [],
        "输出格式": ["txt", "markdown", "html"],
        "库状态": {
            "MarkItDown": MARKITDOWN_AVAILABLE,
            "PyPandoc": PYPANDOC_AVAILABLE,
            "python-docx": PYTHON_DOCX_AVAILABLE,
            "pdfplumber": PDFPLUMBER_AVAILABLE,
            "markdown": MARKDOWN_AVAILABLE
        }
    }
    
    # 根据可用库添加支持的格式
    if MARKITDOWN_AVAILABLE:
        formats["输入格式"].extend(["pdf", "docx", "pptx", "xlsx", "txt", "md", "html"])
    
    if PYTHON_DOCX_AVAILABLE:
        formats["输出格式"].append("docx")
    
    if PYPANDOC_AVAILABLE:
        formats["输出格式"].append("pdf")
    
    return json.dumps(formats, ensure_ascii=False, indent=2)

@mcp.tool()
@performance_monitor
def get_file_info(file_path: str) -> str:
    """
    获取文件信息
    
    Args:
        file_path: 文件路径
    
    Returns:
        文件信息的JSON字符串
    """
    try:
        if not os.path.exists(file_path):
            return json.dumps({"error": "文件不存在"}, ensure_ascii=False)
        
        file_stat = os.stat(file_path)
        file_info = {
            "文件名": os.path.basename(file_path),
            "文件路径": os.path.abspath(file_path),
            "文件大小": f"{file_stat.st_size / 1024:.2f} KB",
            "文件格式": Path(file_path).suffix.lower(),
            "创建时间": time.ctime(file_stat.st_ctime),
            "修改时间": time.ctime(file_stat.st_mtime),
            "是否可读": os.access(file_path, os.R_OK),
            "大小限制检查": config.validate_file_size(file_path)
        }
        
        return json.dumps(file_info, ensure_ascii=False, indent=2)
        
    except Exception as e:
        return json.dumps({"error": f"获取文件信息失败: {str(e)}"}, ensure_ascii=False)

@mcp.tool()
@performance_monitor
def batch_convert(input_dir: str, output_dir: str, target_format: str, file_pattern: str = "*") -> str:
    """
    批量转换文档
    
    Args:
        input_dir: 输入目录
        output_dir: 输出目录
        target_format: 目标格式
        file_pattern: 文件匹配模式（默认为所有文件）
    
    Returns:
        批量转换结果
    """
    try:
        if not os.path.exists(input_dir):
            return json.dumps({"error": "输入目录不存在"}, ensure_ascii=False)
        
        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)
        
        # 查找匹配的文件
        input_path = Path(input_dir)
        files = list(input_path.glob(file_pattern))
        
        results = {
            "总文件数": len(files),
            "成功转换": 0,
            "转换失败": 0,
            "详细结果": []
        }
        
        for file_path in files:
            if file_path.is_file():
                output_file = Path(output_dir) / f"{file_path.stem}.{target_format}"
                result = convert_document(str(file_path), target_format, str(output_file))
                
                if "转换成功" in result:
                    results["成功转换"] += 1
                    status = "成功"
                else:
                    results["转换失败"] += 1
                    status = "失败"
                
                results["详细结果"].append({
                    "文件": str(file_path),
                    "状态": status,
                    "结果": result
                })
        
        return json.dumps(results, ensure_ascii=False, indent=2)
        
    except Exception as e:
        return json.dumps({"error": f"批量转换失败: {str(e)}"}, ensure_ascii=False)

@mcp.tool()
def health_check() -> str:
    """
    服务器健康检查
    
    Returns:
        服务器状态信息
    """
    status = {
        "服务器状态": "运行中",
        "版本": "1.0.0",
        "依赖库状态": {
            "MarkItDown": "可用" if MARKITDOWN_AVAILABLE else "不可用",
            "PyPandoc": "可用" if PYPANDOC_AVAILABLE else "不可用",
            "python-docx": "可用" if PYTHON_DOCX_AVAILABLE else "不可用",
            "pdfplumber": "可用" if PDFPLUMBER_AVAILABLE else "不可用",
            "markdown": "可用" if MARKDOWN_AVAILABLE else "不可用"
        },
        "配置信息": {
            "最大文件大小": f"{config.max_file_size / 1024 / 1024:.1f}MB",
            "临时目录": config.temp_dir,
            "日志级别": config.log_level,
            "启用Pandoc": config.enable_pandoc
        },
        "支持格式": {
            "输入": ["pdf", "docx", "pptx", "xlsx", "txt", "md", "html"] if MARKITDOWN_AVAILABLE else ["txt", "md"],
            "输出": ["txt", "markdown", "html", "docx", "pdf"] if all([PYTHON_DOCX_AVAILABLE, PYPANDOC_AVAILABLE]) else ["txt", "markdown", "html"]
        }
    }
    
    return json.dumps(status, ensure_ascii=False, indent=2)

def main():
    """主函数"""
    logger.info("启动文档转换 MCP 服务器...")
    mcp.run()

if __name__ == "__main__":
    main()